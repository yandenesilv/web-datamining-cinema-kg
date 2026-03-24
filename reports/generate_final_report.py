"""
Generate the final DOCX report for Web Datamining project.
Style matches Cours5/rapport_cours5_KGE.docx.
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Helper functions --
def add_title_page():
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ESILV - A4 S8")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Web Mining & Semantics")
    run.font.size = Pt(16)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Final Project Report")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Cinema Knowledge Graph Pipeline")
    run.font.size = Pt(14)
    run.font.italic = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Yanis DENIZOT")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("March 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')


def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    p.paragraph_format.left_indent = Cm(1)


def add_screenshot_placeholder(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[INSERT SCREENSHOT HERE: {caption}]")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Figure: {caption}")
    run.font.size = Pt(9)
    run.font.italic = True
    doc.add_paragraph()


# ============================================================
# TITLE PAGE
# ============================================================
add_title_page()

# ============================================================
# 1. DATA ACQUISITION & INFORMATION EXTRACTION
# ============================================================
doc.add_heading('1. Data Acquisition & Information Extraction', level=1)

doc.add_heading('1.1 Domain & Seed URLs', level=2)
doc.add_paragraph(
    'We chose the cinema domain as our focus area. Wikipedia provides rich, structured content '
    'about films, directors, actors, awards, and festivals, making it ideal for knowledge extraction.'
)
doc.add_paragraph('Our crawler starts from 9 seed URLs:')
add_bullet('Award pages: Academy Award for Best Picture, Palme d\'Or')
add_bullet('Director pages: Christopher Nolan, Martin Scorsese, Quentin Tarantino')
add_bullet('Actor pages: Cate Blanchett')
add_bullet('Film pages: Parasite, Inception, The Godfather')

doc.add_heading('1.2 Crawler Design & Ethics', level=2)
doc.add_paragraph(
    'The crawler (src/crawl/crawler.py) uses Breadth-First Search (BFS) to explore Wikipedia, '
    'following internal links up to depth 2 from the seed URLs. '
    'Text extraction is handled by Trafilatura, which cleanly removes navigation and HTML boilerplate.'
)
doc.add_paragraph('Ethical considerations implemented:')
add_bullet('robots.txt compliance: checks robots.txt before each request')
add_bullet('Rate limiting: 1-second delay between requests (POLITENESS_DELAY = 1.0)')
add_bullet('User-Agent: identifies as "ESILV-WebMining-Bot/1.0 (student project)"')
add_bullet('Page cap: maximum 150 pages to avoid server overload')
add_bullet('Content deduplication: MD5 hashing prevents re-processing identical pages')
add_bullet('Minimum quality: only pages with 500+ words are kept')

doc.add_heading('1.3 Cleaning & NER Pipeline', level=2)
doc.add_paragraph(
    'The extraction pipeline (src/ie/extractor.py) uses spaCy (en_core_web_lg) for Named Entity Recognition. '
    'We target 5 entity types: PERSON, ORG, GPE, DATE, WORK_OF_ART.'
)
doc.add_paragraph('Relation extraction uses a 3-level strategy:')
add_bullet('Level 1 - Dependency parsing: walk the dependency tree to find verbs governing both entities')
add_bullet('Level 2 - Verb-between: find verbs positioned between two entities in the sentence')
add_bullet('Level 3 - Fallback: "co-occurs with" for entities in the same sentence without a clear verb')

doc.add_paragraph('NER examples from our corpus:')
add_table(
    ['Entity', 'Type', 'Source'],
    [
        ['Christopher Nolan', 'PERSON', 'Inception article'],
        ['Academy Award', 'ORG', 'Award pages'],
        ['United States', 'GPE', 'Multiple pages'],
        ['2019', 'DATE', 'Parasite article'],
        ['The Godfather', 'WORK_OF_ART', 'Film articles'],
    ]
)

doc.add_heading('1.4 Ambiguity Cases', level=2)

doc.add_paragraph().add_run('Case 1: Entity boundary ambiguity').bold = True
doc.add_paragraph(
    '"Francis Ford Coppola directed The Godfather Part II" — spaCy recognizes '
    '"Francis Ford Coppola" as a single PERSON entity, but "The Godfather Part II" may be split '
    'into "The Godfather" (WORK_OF_ART) and "Part II" separately. Our filtering removes fragments '
    'shorter than 2 characters.'
)

doc.add_paragraph().add_run('Case 2: Relation ambiguity').bold = True
doc.add_paragraph(
    '"Nolan\'s Inception starred Leonardo DiCaprio" — the possessive "Nolan\'s" creates ambiguity: '
    'is Nolan the director or producer? Our dependency parser extracts "starred" as the relation '
    'between Inception and DiCaprio, but misses the Nolan-Inception "directed" relation expressed via possessive.'
)

doc.add_paragraph().add_run('Case 3: Type ambiguity').bold = True
doc.add_paragraph(
    '"Columbia Pictures released the film in the United States" — "Columbia Pictures" is recognized '
    'as ORG, but could be classified as a production company (a more specific type in our ontology). '
    'Generic NER types don\'t capture domain-specific subtypes.'
)

# ============================================================
# 2. KB CONSTRUCTION & ALIGNMENT
# ============================================================
doc.add_page_break()
doc.add_heading('2. KB Construction & Alignment', level=1)

doc.add_heading('2.1 RDF Modeling Choices', level=2)
doc.add_paragraph(
    'The knowledge graph uses a private namespace (http://cinema-kb.org/) with an OWL ontology.'
)
doc.add_paragraph('Class hierarchy:')
add_bullet('Thing > Person > Director, Actor')
add_bullet('Thing > CreativeWork > Film')
add_bullet('Thing > Organization > Festival')
add_bullet('Thing > Award')

doc.add_paragraph('Properties with domain/range constraints:')
add_table(
    ['Property', 'Domain', 'Range'],
    [
        ['directedBy', 'Film', 'Person'],
        ['starring', 'Film', 'Person'],
        ['wonAward', 'Film', 'Award'],
        ['producedBy', 'Film', 'Organization'],
        ['releasedIn', 'Film', 'GPE'],
    ]
)
doc.add_paragraph(
    'The kb_builder.py script normalizes 60+ raw relation strings (e.g., "directed by", "direct", "directed") '
    'into a canonical set of 15 predicates using a manually curated mapping.'
)

doc.add_heading('2.2 Entity Linking', level=2)
doc.add_paragraph(
    'Entity linking (src/kg/entity_linker.py) maps the top 200 most frequent private entities '
    'to Wikidata via the REST API. String similarity is computed with SequenceMatcher (threshold: 0.7). '
    'Result: 183 entities successfully linked with owl:sameAs.'
)

doc.add_heading('2.3 Predicate Alignment', level=2)
doc.add_paragraph('17 private predicates were mapped to Wikidata properties using owl:equivalentProperty:')
add_table(
    ['Private Predicate', 'Wikidata Property', 'Description'],
    [
        ['directedBy', 'P57', 'director'],
        ['starring', 'P161', 'cast member'],
        ['wonAward', 'P166', 'award received'],
        ['genre', 'P136', 'genre'],
        ['countryOfOrigin', 'P495', 'country of origin'],
    ]
)

doc.add_heading('2.4 SPARQL Expansion Strategy', level=2)
doc.add_paragraph('The KB expansion (src/kg/kb_expander.py) runs 5 SPARQL queries against the Wikidata endpoint:')
add_bullet('1-hop expansion for all aligned entities')
add_bullet('Award-winning films with full metadata')
add_bullet('Films by aligned directors with cast, genre, country')
add_bullet('Festival nominees (Oscar, Cannes, BAFTA, Golden Globe)')
add_bullet('Director details (birth, death, awards, nationality)')

doc.add_heading('2.5 Final KB Statistics', level=2)
add_table(
    ['Metric', 'Before Expansion', 'After Expansion'],
    [
        ['Total triples', '~10,000', '72,409'],
        ['Total entities', '~3,000', '20,158'],
        ['Distinct relations', '15', '39'],
        ['owl:sameAs links', '0', '183'],
    ]
)
doc.add_paragraph(
    'Top connected entities: United States (1,669 connections), drama film (1,192), GB (471), '
    'National Board of Review: Top Ten Films (464).'
)

# ============================================================
# 3. REASONING (SWRL)
# ============================================================
doc.add_page_break()
doc.add_heading('3. Reasoning (SWRL)', level=1)

doc.add_heading('3.1 SWRL on Family Ontology', level=2)
doc.add_paragraph(
    'We created a family ontology (kg_artifacts/family.owl) with 3 generations of the "Dupont" family '
    'using OWLReady2, and defined 3 SWRL rules:'
)
doc.add_paragraph().add_run('Rule 1 — hasUncle:').bold = True
add_code('Person(?x), hasParent(?x, ?y), hasBrother(?y, ?z) -> hasUncle(?x, ?z)')
doc.add_paragraph().add_run('Rule 2 — hasGrandparent:').bold = True
add_code('Person(?x), hasParent(?x, ?y), hasParent(?y, ?z) -> hasGrandparent(?x, ?z)')
doc.add_paragraph().add_run('Rule 3 — hasCousin:').bold = True
add_code('Person(?x), hasParent(?x, ?y), hasSibling(?y, ?z), hasChild(?z, ?w) -> hasCousin(?x, ?w)')

doc.add_paragraph('Results (all correct):')
add_bullet('Lucas and Emma have uncle Paul; Hugo has uncle Pierre')
add_bullet('Lucas, Emma, Hugo all have grandparents Jean and Marie')
add_bullet('Lucas/Emma are cousins of Hugo (and vice versa)')

add_screenshot_placeholder('SWRL family.owl reasoning output (terminal)')

doc.add_heading('3.2 SWRL on Cinema KB', level=2)
doc.add_paragraph('On the cinema KB, we defined:')
add_code('Film(?f), directedBy(?f, ?p), wonAward(?f, ?a) -> AwardWinningDirector(?p)')
doc.add_paragraph(
    'Inferred: Coppola (The Godfather won Oscar), Bong Joon-ho (Parasite won Oscar + Palme d\'Or), '
    'Tarantino (Pulp Fiction won Palme d\'Or) are AwardWinningDirectors. '
    'Nolan is correctly NOT inferred (Inception did not win Best Picture in our data).'
)

# ============================================================
# 4. KNOWLEDGE GRAPH EMBEDDINGS
# ============================================================
doc.add_page_break()
doc.add_heading('4. Knowledge Graph Embeddings', level=1)

doc.add_heading('4.1 Data Cleaning & Splits', level=2)
doc.add_paragraph(
    'From the 72,409 expanded triples, we kept only entity-relation-entity triples '
    '(excluding literals and metadata), producing 55,294 unique triples.'
)
add_table(
    ['Split', 'Triples', 'Percentage'],
    [
        ['Train', '44,444', '80%'],
        ['Validation', '5,425', '10%'],
        ['Test', '5,425', '10%'],
    ]
)
doc.add_paragraph(
    'Entities appearing in only 1-2 triples were forced into the training set to prevent data leakage.'
)

doc.add_heading('4.2 Hyperparameters', level=2)
doc.add_paragraph('Both models trained with identical hyperparameters for fair comparison:')
add_table(
    ['Parameter', 'Value'],
    [
        ['Embedding dimension', '100'],
        ['Learning rate', '0.001'],
        ['Batch size', '256'],
        ['Epochs', '100'],
        ['Negative samples', '10 per positive'],
        ['Optimizer', 'Adam'],
        ['Device', 'CPU'],
    ]
)

doc.add_heading('4.3 Model Comparison', level=2)
add_table(
    ['Metric', 'TransE', 'ComplEx', 'Winner'],
    [
        ['MRR', '0.133', '0.009', 'TransE'],
        ['Hits@1', '0.067', '0.003', 'TransE'],
        ['Hits@3', '0.161', '0.007', 'TransE'],
        ['Hits@10', '0.251', '0.017', 'TransE'],
    ]
)
doc.add_paragraph(
    'TransE significantly outperforms ComplEx. This is expected: our KB has mostly functional relations '
    '(1-to-1, 1-to-N) which suit TransE\'s translational assumption (h + r = t). ComplEx needs more '
    'epochs and hyperparameter tuning.'
)

doc.add_heading('4.4 Head vs Tail Prediction (TransE)', level=3)
add_table(
    ['Metric', 'Head', 'Tail', 'Both'],
    [
        ['MRR', '0.110', '0.157', '0.133'],
        ['Hits@10', '20.0%', '30.2%', '25.1%'],
    ]
)
doc.add_paragraph(
    'Tail prediction outperforms head prediction: the model is better at predicting '
    '"what object relates to this subject" than the reverse.'
)

doc.add_heading('4.5 KB Size Sensitivity', level=2)
add_table(
    ['Size', 'Triples', 'Entities', 'MRR', 'Hits@10'],
    [
        ['20K', '20,000', '11,441', '0.042', '0.111'],
        ['Full', '55,294', '19,733', '0.137', '0.257'],
    ]
)
doc.add_paragraph(
    'The full dataset achieves 3.3x better MRR than the 20K subset, confirming that '
    'larger KBs improve embedding stability.'
)

doc.add_heading('4.6 t-SNE Visualization', level=2)
doc.add_paragraph(
    'We applied t-SNE to the top 3,000 most connected entity embeddings, colored by ontology class. '
    'While semantic groupings are visible (persons cluster, places form distinct groups), '
    'the clustering does not cleanly separate ontology classes because 82% of entities lack explicit rdf:type.'
)
add_screenshot_placeholder('t-SNE clustering visualization (figures/tsne_clustering.png)')

# ============================================================
# 5. RAG OVER RDF/SPARQL
# ============================================================
doc.add_page_break()
doc.add_heading('5. RAG over RDF/SPARQL', level=1)

doc.add_heading('5.1 Setup', level=2)
doc.add_paragraph(
    'Machine: Windows 11, CPU (no GPU required). '
    'Model: qwen2.5:3b via Ollama (localhost:11434). '
    'Graph: expanded_kb.ttl (83,065 triples including ontology metadata).'
)

doc.add_heading('5.2 Schema Summary', level=2)
doc.add_paragraph(
    'The RAG pipeline loads the KB and builds a compact schema summary for the LLM prompt, containing: '
    'registered prefixes, up to 30 distinct predicates, up to 15 classes, and 10 sample triples. '
    'This provides the LLM with enough context to generate valid SPARQL.'
)

doc.add_heading('5.3 NL to SPARQL Prompt', level=2)
doc.add_paragraph(
    'We use a few-shot prompt with concrete SPARQL examples matching our KB structure. '
    'Key predicates are listed explicitly (P57=director, P161=cast, P166=award, etc.) '
    'and 3 example question-SPARQL pairs are provided.'
)

doc.add_heading('5.4 Self-Repair Mechanism', level=2)
doc.add_paragraph('When the generated SPARQL fails execution, the pipeline:')
add_bullet('1. Captures the error message')
add_bullet('2. Sends a repair prompt with schema + question + failed query + error')
add_bullet('3. Extracts and re-executes the corrected SPARQL')
add_bullet('4. Retries up to 2 times before giving up')

doc.add_heading('5.5 Evaluation', level=2)

doc.add_paragraph().add_run('Automated evaluation (LLM-generated SPARQL, qwen2.5:3b):').bold = True
add_table(
    ['#', 'Question', 'Baseline', 'RAG (auto)', 'Repair?'],
    [
        ['1', 'Who directed The Godfather?', 'Correct', '0 rows', 'No'],
        ['2', 'What awards did Parasite receive?', 'Correct', '0 rows', 'No'],
        ['3', 'Films by Martin Scorsese', 'Correct', '0 rows', 'No'],
        ['4', 'Genre of Inception?', 'Correct', '0 rows', 'No'],
        ['5', 'US directors?', 'Correct', '0 rows', 'No'],
        ['6', 'How many films in KB?', 'Cannot answer', '1 row (1678)', 'No'],
        ['7', 'Best Picture winners?', 'Correct', '0 rows', 'No'],
    ]
)
doc.add_paragraph(
    'Result: 1/7 success with LLM-generated SPARQL. The 3b model struggles with complex Wikidata URIs. '
    'Q6 succeeded because the COUNT pattern was provided as a few-shot example.'
)

doc.add_paragraph().add_run('Manual SPARQL verification (hand-written queries):').bold = True
add_table(
    ['#', 'Question', 'SPARQL Result', 'Correct?'],
    [
        ['1', 'Who directed The Godfather?', 'Francis Ford Coppola (Q56094)', 'Yes'],
        ['2', 'Awards for Parasite?', 'Academy Award Best Picture, Palme d\'Or + 8 more', 'Yes'],
        ['3', 'Films by Scorsese?', 'The Departed, Wolf of Wall Street + 8 more', 'Yes'],
        ['4', 'How many films with directors?', '1,678 films', 'Yes'],
        ['5', 'Genre of Inception?', 'No genre data in KB', 'N/A'],
    ]
)
doc.add_paragraph(
    'Result: 4/5 success with hand-written SPARQL. The KB itself is correct and rich.'
)

doc.add_paragraph().add_run('Key observations:').bold = True
add_bullet('The KB contains rich, queryable data (83,065 triples, 1,678 films with directors)')
add_bullet('Small LLMs (0.5b-3b) cannot reliably generate SPARQL with complex Wikidata URIs')
add_bullet('The self-repair mechanism correctly detects failures (activated 4/7 times with 0.5b model)')
add_bullet('Larger models (7b+) would likely perform significantly better')
add_bullet('Baseline provides fluent but sometimes hallucinated answers')
add_bullet('RAG provides grounded, verifiable answers when SPARQL is correct')

doc.add_heading('5.6 Demo', level=2)
doc.add_paragraph(
    'The CLI demo (src/rag/lab_rag_sparql_gen.py) provides an interactive interface. '
    'Below is an example showing baseline vs RAG for "Who directed The Godfather?":'
)
add_screenshot_placeholder('RAG CLI demo showing baseline vs SPARQL results (open reports/rag_screenshot.html in browser)')

# ============================================================
# 6. CRITICAL REFLECTION
# ============================================================
doc.add_page_break()
doc.add_heading('6. Critical Reflection', level=1)

doc.add_heading('6.1 KB Quality Impact', level=2)
doc.add_paragraph(
    'The dual-namespace issue (private predicates + Wikidata properties) creates signal splitting '
    'in KGE training. P57 and directedBy encode the same meaning but are treated as different relations. '
    'Physically merging equivalent predicates before training would produce stronger embeddings.'
)

doc.add_heading('6.2 Noise from Expansion', level=2)
doc.add_paragraph(
    'P161 (cast member) accounts for 47% of all triples (25,783/55,294), biasing embeddings '
    'toward cast-related patterns. Hub entities like "United States" (1,669 connections) distort '
    'the embedding space. A capped expansion strategy would produce more balanced distributions.'
)

doc.add_heading('6.3 Rule-based vs Embedding-based Reasoning', level=2)
add_table(
    ['Aspect', 'SWRL (Rules)', 'KGE (Embeddings)'],
    [
        ['Nature', 'Exact, deterministic', 'Approximate, probabilistic'],
        ['Inference', 'Guaranteed if premises hold', 'Ranks by plausibility'],
        ['Discovery', 'Cannot find new patterns', 'Can suggest missing links'],
        ['Scalability', 'Manual rule creation', 'Learns from data'],
        ['Explainability', 'Fully transparent', 'Black box'],
    ]
)
doc.add_paragraph('The two approaches are complementary: SWRL for strict business rules, KGE for knowledge discovery.')

doc.add_heading('6.4 Open-World vs Closed-World Assumption', level=2)
doc.add_paragraph(
    'SWRL/OWL operates under the Open-World Assumption (OWA): absence of a triple does not imply '
    'it\'s false. KGE training uses the Closed-World Assumption (CWA): missing triples are negative '
    'examples. This mismatch explains modest MRR scores (0.133).'
)

doc.add_heading('6.5 RAG Limitations & Small LLM Analysis', level=2)
doc.add_paragraph(
    'Our RAG evaluation revealed a critical finding: small LLMs (0.5b-3b parameters) lack '
    'the ability to generate syntactically correct SPARQL with complex Wikidata URIs. '
    'The few-shot prompting improved results (Q6 succeeded), but most queries still fail. '
    'This highlights the importance of model size for structured output generation tasks.'
)

doc.add_heading('6.6 What We Would Improve', level=2)
add_bullet('Predicate merging: physically unify equivalent predicates before KGE training')
add_bullet('Balanced expansion: cap triples per relation type to reduce P161 dominance')
add_bullet('Hub filtering: remove or downsample generic hub entities')
add_bullet('Larger LLM: use 7b+ model for SPARQL generation, or fine-tune on SPARQL examples')
add_bullet('Template fallback: add rule-based SPARQL templates as hybrid approach')
add_bullet('Type-aware embeddings: include rdf:type as regular relations for better clustering')

# ============================================================
# SAVE
# ============================================================
output_path = 'reports/final_report.docx'
doc.save(output_path)
print(f'Report saved to {output_path}')
print(f'Total pages: ~{len(doc.paragraphs) // 30} (estimated)')
